import difflib
import logging
import re
import docx
import difflib
from typing import List, Dict, Any


"""Word 文檔處理模組"""
class WordProcessor:
    """Word 文檔處理類別"""

    def __init__(self):
        """初始化 Word 處理器"""
        self.logger = logging.getLogger(self.__class__.__name__)
        self.document = None
        self.text_content = ""  # 確保這行存在且初始化為空字符串
        self.paragraphs = []
        self.processed_paragraphs = []
        self.word_file_path = None
        self.edited_paragraphs = {} # 儲存被編輯過的段落 {index: edited_text}

    def load_document(self, file_path: str) -> bool:
        """
        載入 Word 文檔
        :param file_path: Word 文檔路徑
        :return: 是否成功載入
        """
        try:
            self.document = docx.Document(file_path)
            self.word_file_path = file_path
            self.extract_text()
            self.preprocess_paragraphs()  # 添加這行以進一步處理段落
            self.logger.info(f"成功載入 Word 文檔: {file_path}, 共 {len(self.paragraphs)} 個段落")
            return True
        except Exception as e:
            self.logger.error(f"載入 Word 文檔失敗: {e}")
            return False

    def extract_text(self) -> None:
        """
        提取 Word 文檔中的純文本，每一段作為一個獨立條目
        """
        if not self.document:
            return

        self.paragraphs = []
        self.processed_paragraphs = []

        # 遍歷文檔中的每個段落
        for para in self.document.paragraphs:
            text = para.text.strip()
            # 只保留非空段落
            if text:
                self.paragraphs.append(text)  # 保留原始文本
                # 處理文本：忽略標點符號和空格（用於比對）
                processed_text = self._remove_punctuation_and_spaces(text)
                self.processed_paragraphs.append(processed_text)

        # 將所有段落合併為一個完整文本，用於其他地方可能需要的情況
        self.text_content = "\n".join(self.paragraphs)

    def _remove_punctuation_and_spaces(self, text: str) -> str:
        """
        移除文本中的標點符號和空格（用於比對）
        :param text: 原始文本
        :return: 處理後的文本
        """
        # 移除所有標點符號和空格
        return re.sub(r'[^\w\s]|[\s]', '', text)

    def compare_with_srt(self, srt_texts: List[str]) -> Dict[int, Dict[str, Any]]:
        """
        比較 Word 文檔與 SRT 文本，一一對應
        :param srt_texts: SRT 字幕文本列表
        :return: 比對結果字典 {srt_index: {'match': bool, 'word_text': str, 'word_index': int, 'difference': str}}
        """
        if not self.paragraphs:
            return {}

        comparison_results = {}

        # 確認 Word 段落數與 SRT 文本數的對應關係
        for i, srt_text in enumerate(srt_texts):
            word_index = i  # 默認情況下，SRT 索引對應相同的 Word 段落索引

            # 處理 SRT 文本，移除標點和空格以便比較
            processed_srt = self._remove_punctuation_and_spaces(srt_text)

            # 檢查是否存在對應的 Word 段落
            if word_index < len(self.paragraphs):
                word_text = self.paragraphs[word_index]
                processed_word = self.processed_paragraphs[word_index]

                # 計算相似度
                similarity = self._calculate_similarity(processed_srt, processed_word)

                # 如果相似度足夠高，則認為匹配（可調整閾值）
                if similarity > 0.7:
                    comparison_results[i] = {
                        'match': True,
                        'word_text': word_text,
                        'word_index': word_index,
                        'difference': ""
                    }
                else:
                    # 生成差異信息
                    diff_text = self._generate_difference_details(srt_text, word_text)

                    comparison_results[i] = {
                        'match': False,
                        'word_text': word_text,
                        'word_index': word_index,
                        'difference': diff_text
                    }
            else:
                # SRT 條目超出了 Word 段落數量
                comparison_results[i] = {
                    'match': False,
                    'word_text': "Word 文檔中無對應段落",
                    'word_index': -1,
                    'difference': "超出 Word 文檔範圍"
                }

        return comparison_results

    def preprocess_paragraphs(self):
        """進一步處理段落，確保一行文本對應一個要點"""
        processed_paras = []

        for para in self.paragraphs:
            # 如果段落包含換行符，可能需要拆分
            if '\n' in para:
                lines = [line.strip() for line in para.split('\n') if line.strip()]
                processed_paras.extend(lines)
            else:
                processed_paras.append(para)

        # 更新段落列表
        self.paragraphs = processed_paras
        # 更新處理後的段落列表
        self.processed_paragraphs = [self._remove_punctuation_and_spaces(p) for p in processed_paras]

    def _generate_difference_details(self, processed_srt: str, processed_word: str) -> str:
        """
        生成差異詳情
        :param processed_srt: 處理後的 SRT 文本
        :param processed_word: 處理後的 Word 文本
        :return: 差異詳情
        """
        # 使用 difflib 找出差異
        d = difflib.Differ()
        diff = list(d.compare(processed_srt, processed_word))

        srt_only = ''.join([c[2] for c in diff if c.startswith('- ')])
        word_only = ''.join([c[2] for c in diff if c.startswith('+ ')])

        result = []
        if srt_only:
            result.append(f"SRT={srt_only}")
        if word_only:
            result.append(f"WORD={word_only}")

        return "｜".join(result) if result else "內容不完全匹配"

    def _calculate_similarity(self, text1: str, text2: str) -> float:
        """
        計算兩個文本之間的相似度
        :param text1: 第一個文本
        :param text2: 第二個文本
        :return: 相似度分數 (0-1)
        """
        # 改用更準確的相似度算法

        # 1. 使用 Jaccard 相似度計算詞集合的相似度
        set1 = set(text1.split())
        set2 = set(text2.split())

        if not set1 or not set2:
            return 0

        jaccard_similarity = len(set1.intersection(set2)) / len(set1.union(set2))

        # 2. 考慮詞序
        sequence_matcher = difflib.SequenceMatcher(None, text1, text2)
        sequence_similarity = sequence_matcher.ratio()

        # 3. 綜合考慮詞集合和詞序的相似度
        combined_similarity = (jaccard_similarity * 0.5) + (sequence_similarity * 0.5)

        return combined_similarity

    def edit_paragraph(self, index: int, new_text: str) -> bool:
        """
        編輯 Word 段落
        :param index: 段落索引
        :param new_text: 新文本
        :return: 是否成功編輯
        """
        if 0 <= index < len(self.paragraphs):
            # 保存編輯後的文本
            self.edited_paragraphs[index] = new_text
            # 更新處理後的文本（用於比對）
            self.processed_paragraphs[index] = self._remove_punctuation_and_spaces(new_text)
            self.logger.info(f"編輯 Word 段落 {index}: {new_text}")
            return True
        return False

    def get_paragraph_text(self, index: int) -> str:
        """
        獲取段落文本
        :param index: 段落索引
        :return: 段落文本
        """
        # 優先返回編輯後的文本
        if index in self.edited_paragraphs:
            return self.edited_paragraphs[index]

        if 0 <= index < len(self.paragraphs):
            return self.paragraphs[index]
        return ""

    def get_paragraphs_count(self) -> int:
        """
        獲取段落數量
        :return: 段落數量
        """
        return len(self.paragraphs)

    def split_paragraph(self, index, texts):
        """
        將一個段落分割為多個
        :param index: 原始段落的索引
        :param texts: 分割後的文本列表
        :return: 新段落的索引列表
        """
        if index < 0 or index >= len(self.paragraphs):
            self.logger.error(f"無效的段落索引: {index}")
            return []

        try:
            # 記錄原始段落內容（用於調試）
            original_text = self.paragraphs[index]
            self.logger.debug(f"分割段落 {index}: {original_text} => {len(texts)} 個新段落")

            # 保存原始索引的數據用於後續調整
            affected_indices = {}

            # 記錄所有索引大於等於index的項目，因為它們的位置可能會發生變化
            for i in range(index, len(self.paragraphs)):
                affected_indices[i] = {
                    'text': self.paragraphs[i],
                    'processed': self.processed_paragraphs[i],
                    'edited': i in self.edited_paragraphs
                }
                if i in self.edited_paragraphs:
                    affected_indices[i]['edited_text'] = self.edited_paragraphs[i]

            # 刪除原始段落
            self.paragraphs.pop(index)
            self.processed_paragraphs.pop(index)

            # 如果原段落有編輯狀態，記錄並移除它
            had_edit = index in self.edited_paragraphs
            if had_edit:
                edited_text = self.edited_paragraphs[index]
                del self.edited_paragraphs[index]

            # 插入新段落
            new_indices = []
            for i, text in enumerate(texts):
                insert_index = index + i
                self.paragraphs.insert(insert_index, text)
                self.processed_paragraphs.insert(insert_index, self._remove_punctuation_and_spaces(text))
                new_indices.append(insert_index)

                # 如果原段落有編輯狀態，將第一個新段落設為編輯狀態
                if had_edit and i == 0:
                    self.edited_paragraphs[insert_index] = text

            # 調整所有受影響索引的位置
            offset = len(texts) - 1  # 計算索引偏移量

            # 處理所有受影響的索引
            for old_idx, data in sorted(affected_indices.items(), reverse=True):
                # 跳過已處理的原始索引
                if old_idx == index:
                    continue

                # 計算新的索引位置
                new_idx = old_idx + offset

                # 如果此段落有編輯標記，更新編輯標記的索引
                if data['edited'] and old_idx in self.edited_paragraphs:
                    self.edited_paragraphs[new_idx] = self.edited_paragraphs[old_idx]
                    del self.edited_paragraphs[old_idx]

            # 更新文本內容
            self.text_content = "\n".join(self.paragraphs)
            self.logger.debug(f"段落分割完成，新段落索引: {new_indices}")

            return new_indices

        except Exception as e:
            self.logger.error(f"分割段落時出錯: {e}")
            return []